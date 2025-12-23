"""
TailorMatch - Diploma Thesis Generator (Part 1: Setup and Title)
Run this in Google Colab to generate complete diploma thesis
"""

# Install required library
!pip install python-docx -q

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from google.colab import files
from datetime import datetime

# Create document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

# Helper functions
def add_heading_centered(doc, text, level=0):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return heading

def add_paragraph_justified(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_page_break(doc):
    doc.add_page_break()

# ==================== TITLE PAGE ====================
# Study program info (Latvian)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Studiju programma\t\t47483 Datorsistēmās")
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

# Department
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Dabaszinātņu un datortehnoloģiju katedra")
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

# Author name
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Rustamjon Abduvahobov")
run.font.size = Pt(14)
run.bold = True

doc.add_paragraph()
doc.add_paragraph()

# Thesis title (Latvian)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('"Mobilās Lietotnes Izstrāde Šuvēju un Klientu Savienošanai, Izmantojot Flutter un Firebase Tehnoloģijas"')
run.font.size = Pt(14)
run.bold = True

doc.add_paragraph()
doc.add_paragraph()

# Type
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MAĢISTRA DARBS")
run.font.size = Pt(16)
run.bold = True

for _ in range(8):
    doc.add_paragraph()

# City and year
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("RĪGA 2026")
run.font.size = Pt(14)

add_page_break(doc)

# ==================== TITLE PAGE (English) ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Study program\t\t47483 Computer System")
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Natural Sciences and Computer Technologies Department")
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Rustamjon Abduvahobov")
run.font.size = Pt(14)
run.bold = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('"Development of a Mobile Application for Connecting Tailors and Clients Using Flutter and Firebase Technologies"')
run.font.size = Pt(14)
run.bold = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MASTER THESIS")
run.font.size = Pt(16)
run.bold = True

for _ in range(8):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("RIGA 2026")
run.font.size = Pt(14)

add_page_break(doc)

# ==================== ABSTRACT ====================
add_heading_centered(doc, "Abstract", 1)

abstract_text = """The tailoring industry in Uzbekistan faces significant challenges in connecting skilled tailors with clients seeking custom clothing services. Traditional methods of finding tailors rely heavily on word-of-mouth recommendations and physical proximity, limiting both client access to quality services and tailor exposure to potential customers. This inefficiency results in underutilized tailor capacity and unsatisfied client demand.

The objective of this thesis is to design and implement a cross-platform mobile application that facilitates seamless connections between tailors and clients, enabling service discovery, real-time communication, and order management through a unified digital platform.

The application was developed using the Flutter framework for cross-platform compatibility, with Firebase providing backend services including authentication, real-time database (Cloud Firestore), and file storage. The system architecture follows a service-oriented pattern with clear separation between presentation, business logic, and data layers.

The implemented TailorMatch application successfully demonstrates core functionality including user authentication with role-based access (client/tailor), service listing and search capabilities, real-time chat messaging between clients and tailors, multi-aspect rating system (quality, size, speed, service, price) following Yandex-style evaluation patterns, order management with status tracking across four categories (New, In Progress, Ready, Completed), tailor calendar for busy day management, and full multilingual support (Uzbek, Russian, English).

Performance testing validates that the application maintains responsive user experience with message delivery under 500ms latency. The modular architecture enables straightforward extension for future features including payment integration and geolocation services.

This solution provides a practical foundation for digitizing the tailoring service industry in Uzbekistan, with potential applicability to similar service-matching domains in other regional markets."""

add_paragraph_justified(doc, abstract_text)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Keywords: ").bold = True
p.add_run("mobile application development, Flutter framework, Firebase backend, cross-platform development, real-time messaging, service marketplace, user authentication, Cloud Firestore, rating system, multilingual support")

add_page_break(doc)

# ==================== ANOTĀCIJA (Latvian) ====================
add_heading_centered(doc, "Anotācija", 1)

anotacija_text = """Šūšanas nozare Uzbekistānā saskaras ar būtiskām grūtībām, savienojot kvalificētus šuvējus ar klientiem, kuri meklē individuālus apģērbu pakalpojumus. Tradicionālās metodes šuvēju atrašanai lielā mērā balstās uz ieteikumiem un fizisko tuvumu, ierobežojot gan klientu piekļuvi kvalitatīviem pakalpojumiem, gan šuvēju saskarsmi ar potenciālajiem klientiem.

Šī darba mērķis ir izstrādāt un ieviest starpplatformu mobilo lietotni, kas nodrošina netraucētu savienojumu starp šuvējiem un klientiem, ļaujot atrast pakalpojumus, reāllaika saziņu un pasūtījumu pārvaldību vienotā digitālā platformā.

Lietotne tika izstrādāta, izmantojot Flutter ietvaru starpplatformu saderībai, ar Firebase, kas nodrošina aizmugures pakalpojumus, tostarp autentifikāciju, reāllaika datubāzi (Cloud Firestore) un failu glabātuvi.

Ieviestā TailorMatch lietotne veiksmīgi demonstrē pamata funkcionalitāti, tostarp lietotāju autentifikāciju ar lomu piekļuvi (klients/šuvējs), pakalpojumu sarakstu un meklēšanas iespējas, reāllaika tērzēšanas ziņapmaiņu, daudzaspektu vērtēšanas sistēmu, pasūtījumu pārvaldību ar statusa izsekošanu četrās kategorijās un pilnu daudzvalodu atbalstu (uzbeku, krievu, angļu)."""

add_paragraph_justified(doc, anotacija_text)

add_page_break(doc)

# ==================== TABLE OF CONTENTS ====================
add_heading_centered(doc, "Table of Contents", 1)

toc_items = [
    ("Abstract", "3"),
    ("Anotācija", "4"),
    ("Table of Contents", "5"),
    ("Research Direction", "6"),
    ("Abbreviations", "7"),
    ("Preamble", "9"),
    ("Term of Reference", "10"),
    ("Defended Thesis", "10"),
    ("1. Theoretical Foundations of Mobile Application Development", "11"),
    ("   1.1. Introduction to Cross-Platform Development", "11"),
    ("   1.2. Flutter Framework Architecture", "14"),
    ("   1.3. Firebase Backend Services", "18"),
    ("   1.4. Service Marketplace Platforms Analysis", "22"),
    ("   1.5. Summary", "26"),
    ("2. System Design and Architecture", "27"),
    ("   2.1. Requirements Analysis", "27"),
    ("   2.2. System Architecture Design", "31"),
    ("   2.3. Database Schema Design", "35"),
    ("   2.4. User Interface Design Principles", "39"),
    ("   2.5. Security Architecture", "43"),
    ("   2.6. Summary", "46"),
    ("3. Implementation of TailorMatch Application", "47"),
    ("   3.1. Development Environment Setup", "47"),
    ("   3.2. Authentication Module Implementation", "49"),
    ("   3.3. User Profile and Service Management", "53"),
    ("   3.4. Real-Time Chat System", "57"),
    ("   3.5. Order Management System", "60"),
    ("   3.6. Rating and Review System", "63"),
    ("   3.7. Multilingual Support Implementation", "66"),
    ("   3.8. Summary", "68"),
    ("4. Testing and Evaluation", "69"),
    ("   4.1. Testing Methodology", "69"),
    ("   4.2. Functional Testing Results", "71"),
    ("   4.3. Performance Evaluation", "73"),
    ("   4.4. User Acceptance Testing", "75"),
    ("   4.5. Summary", "77"),
    ("5. Conclusions", "78"),
    ("6. References", "80"),
]

for item, page in toc_items:
    p = doc.add_paragraph()
    if item.startswith("   "):
        p.paragraph_format.left_indent = Inches(0.5)
        p.add_run(item.strip())
    else:
        p.add_run(item).bold = not item.startswith("   ")
    tab_stops = p.paragraph_format.tab_stops
    p.add_run("\t" + page)

add_page_break(doc)

# ==================== RESEARCH DIRECTION ====================
add_heading_centered(doc, "Research Direction", 1)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Topicality").bold = True
p.add_run(" – The digital transformation of service industries has become increasingly critical in developing economies. In Uzbekistan, the tailoring sector remains largely traditional, with limited technological integration. Mobile applications offer unprecedented opportunities to bridge the gap between service providers and consumers, improving accessibility, efficiency, and customer satisfaction. The global mobile application market continues to expand, with service marketplace platforms demonstrating significant growth potential.")

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Core Problem").bold = True
p.add_run(" – The fundamental challenge in the tailoring industry is the inefficient matching of tailors with clients. Clients struggle to discover qualified tailors, assess their work quality, and communicate requirements effectively. Tailors face difficulties in managing orders, scheduling, and building customer relationships. Traditional approaches lack transparency, real-time communication, and systematic feedback mechanisms.")

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Aim").bold = True
p.add_run(" – To design and implement a comprehensive mobile application that addresses the connectivity challenges between tailors and clients by:")

bullets = [
    "Providing a platform for tailor discovery and service presentation",
    "Enabling real-time communication through integrated chat functionality",
    "Implementing order management with status tracking",
    "Developing a multi-aspect rating system for quality assessment",
    "Supporting multilingual interface for diverse user base",
    "Creating scalable architecture for future feature expansion"
]

for bullet in bullets:
    doc.add_paragraph(bullet, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Hypothesis").bold = True
p.add_run(" – A well-designed mobile application utilizing modern cross-platform technologies (Flutter) and cloud backend services (Firebase) can effectively digitize the tailor-client interaction process, resulting in improved service discovery, enhanced communication, and increased customer satisfaction compared to traditional methods.")

add_page_break(doc)

# ==================== ABBREVIATIONS ====================
add_heading_centered(doc, "Abbreviations", 1)

abbreviations = [
    ("API", "Application Programming Interface"),
    ("APK", "Android Package Kit"),
    ("BLoC", "Business Logic Component"),
    ("CLI", "Command Line Interface"),
    ("CRUD", "Create, Read, Update, Delete"),
    ("CSS", "Cascading Style Sheets"),
    ("DAO", "Data Access Object"),
    ("DI", "Dependency Injection"),
    ("DTO", "Data Transfer Object"),
    ("FCM", "Firebase Cloud Messaging"),
    ("GDPR", "General Data Protection Regulation"),
    ("GPS", "Global Positioning System"),
    ("GUI", "Graphical User Interface"),
    ("HTTPS", "Hypertext Transfer Protocol Secure"),
    ("IDE", "Integrated Development Environment"),
    ("iOS", "iPhone Operating System"),
    ("JSON", "JavaScript Object Notation"),
    ("JWT", "JSON Web Token"),
    ("MVVM", "Model-View-ViewModel"),
    ("NoSQL", "Not Only SQL"),
    ("OOP", "Object-Oriented Programming"),
    ("REST", "Representational State Transfer"),
    ("SDK", "Software Development Kit"),
    ("SPA", "Single Page Application"),
    ("SQL", "Structured Query Language"),
    ("SSL", "Secure Sockets Layer"),
    ("TLS", "Transport Layer Security"),
    ("UI", "User Interface"),
    ("UUID", "Universally Unique Identifier"),
    ("UX", "User Experience"),
    ("VM", "Virtual Machine"),
    ("XML", "Extensible Markup Language"),
]

for abbr, full in abbreviations:
    p = doc.add_paragraph()
    p.add_run(abbr).bold = True
    p.add_run(f" – {full}")

add_page_break(doc)

# ==================== PREAMBLE ====================
add_heading_centered(doc, "Preamble", 1)

preamble = """The evolution of mobile technology has fundamentally transformed how services are discovered, consumed, and evaluated. From transportation (Uber, Bolt) to accommodation (Airbnb) to food delivery (Yandex Eats, Glovo), mobile platforms have revolutionized traditional service industries by creating efficient marketplaces that connect providers with consumers.

The tailoring industry, despite being one of the oldest and most essential service sectors, has remained largely untouched by this digital transformation, particularly in Central Asian markets. In Uzbekistan, with its rich textile heritage and significant population of skilled tailors, the disconnect between service providers and potential clients represents both a challenge and an opportunity.

Traditional methods of finding a tailor typically involve personal recommendations, physical visits to multiple establishments, and time-consuming negotiations. This process is inefficient for clients seeking quality services and limits tailors' ability to expand their customer base beyond their immediate geographic vicinity.

The advent of cross-platform mobile development frameworks, particularly Flutter, has democratized mobile application development. Combined with comprehensive backend-as-a-service platforms like Firebase, it is now possible to develop sophisticated, scalable applications with relatively modest resources.

This thesis documents the design, development, and evaluation of TailorMatch, a mobile application specifically created to address the connectivity challenges in the tailoring industry. The application leverages modern technologies to provide a seamless experience for both tailors seeking to showcase their services and clients searching for quality tailoring solutions.

The significance of this work extends beyond the immediate technical implementation. It demonstrates a model that can be adapted for similar service industries in developing markets, contributing to the broader digital transformation of traditional economies."""

add_paragraph_justified(doc, preamble)

add_page_break(doc)

# ==================== TERM OF REFERENCE ====================
add_heading_centered(doc, "Term of Reference", 1)

term_ref = """The tailoring service sector in Uzbekistan faces significant challenges in the digital age. While mobile technology has transformed numerous industries, the connection between skilled tailors and potential clients remains largely dependent on traditional, inefficient methods. This gap presents an opportunity for technological intervention.

The primary aim of this research is to develop a practical mobile application solution that addresses these challenges while demonstrating the effective use of modern cross-platform development technologies. The research encompasses:

1. Analysis of existing service marketplace platforms and their architectural approaches
2. Evaluation of cross-platform mobile development frameworks with emphasis on Flutter
3. Investigation of Firebase backend services for real-time application requirements
4. Design and implementation of a complete mobile application with comprehensive features
5. Testing and validation of the implemented solution

The scope includes both client-facing and tailor-facing functionality, ensuring the platform serves all stakeholders in the tailoring service ecosystem. Special attention is given to multilingual support, reflecting the diverse linguistic landscape of the target market.

The successful completion of this project will result in a functional mobile application that can be deployed for real-world use, demonstrating the viability of the technological approach and providing a foundation for future enhancements."""

add_paragraph_justified(doc, term_ref)

doc.add_paragraph()
doc.add_paragraph()

# ==================== DEFENDED THESIS ====================
add_heading_centered(doc, "Defended Thesis", 1)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Thesis defended is the following:").bold = True

doc.add_paragraph()

defended_items = [
    "Has been offered a comprehensive mobile application architecture utilizing Flutter framework and Firebase backend services that enables efficient connection between tailors and clients in the Uzbekistan market;",
    "Has been modified the traditional service discovery approach by implementing a multi-aspect rating system (quality, size, speed, service, price) following Yandex-style evaluation patterns, providing transparent and detailed feedback mechanisms;",
    "Has been suggested a hybrid solution for real-time communication and order management, combining Cloud Firestore real-time database capabilities with structured order status workflows, enabling seamless interaction between service providers and consumers."
]

for i, item in enumerate(defended_items, 1):
    p = doc.add_paragraph()
    p.add_run(f"{i}. ").bold = True
    p.add_run(item)

add_page_break(doc)

# Save Part 1
doc.save('/content/TailorMatch_Thesis_Part1.docx')
print("✅ Part 1 saved successfully!")
print("Now run Part 2 code...")
