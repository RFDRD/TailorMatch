"""
TailorMatch - Texnik Topshiriq (Technical Specification) Generator
Bu kod Google Colab'da ishlatiladi va Word hujjatini yaratadi
"""

# Kerakli kutubxonalarni o'rnatish
!pip install python-docx -q

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from google.colab import files
from datetime import datetime

# Hujjat yaratish
doc = Document()

# Stillar sozlash
def set_heading_style(paragraph, size=18, bold=True, color=None):
    run = paragraph.runs[0]
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_bullet_point(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

# ==================== TITUL SAHIFASI ====================
title = doc.add_heading('TEXNIK TOPSHIRIQ', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('TAILORMATCH')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(233, 30, 99)  # Pink color

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run("O'zbekistondagi chevarlar va mijozlar uchun\nmobil ilova")
run.font.size = Pt(14)

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run(f'Versiya: 1.0\n')
info.add_run(f'Sana: {datetime.now().strftime("%d.%m.%Y")}\n')
info.add_run('Muallif: Rustamjon Abduvahobov')

doc.add_page_break()

# ==================== MUNDARIJA ====================
doc.add_heading('MUNDARIJA', level=1)
mundarija = [
    "1. Loyihaning vazifalari va maqsadlarini belgilash",
    "2. Loyiha maqsadi",
    "3. Analitik vazifalar",
    "4. Rivojlanish vazifalari",
    "5. Tashkiliy vazifalar",
    "6. Sinov va sifat nazorati",
    "7. Funksional talablar",
    "8. Texnik talablar",
    "9. Dizayn va interfeys",
    "10. Xavfsizlik",
    "11. Muddatlar va bosqichlar"
]
for item in mundarija:
    doc.add_paragraph(item)

doc.add_page_break()

# ==================== 1. LOYIHANING VAZIFALARI ====================
doc.add_heading("1. LOYIHANING VAZIFALARI VA MAQSADLARINI BELGILASH", level=1)

p = doc.add_paragraph()
p.add_run("TailorMatch").bold = True
p.add_run(" — bu O'zbekistondagi chevarlar va mijozlar o'rtasida aloqa o'rnatish uchun mo'ljallangan zamonaviy mobil ilova. Loyiha tikuvchilik sohasini raqamlashtirish va mijozlarga qulay xizmat ko'rsatish maqsadida ishlab chiqilmoqda.")

doc.add_heading("1.1 Asosiy vazifalar", level=2)
vazifalar = [
    "Foydalanuvchilarga chevarlarni qidirish va ko'rish imkoniyatini berish",
    "Onlayn buyurtma berish tizimini joriy etish",
    "Chevarlar uchun buyurtmalarni boshqarish panelini yaratish",
    "Mijoz va chevar o'rtasida samarali chat aloqasini ta'minlash",
    "Viloyat bo'yicha chevarlarni filtrlash imkoniyati",
    "Reyting va baho berish tizimini joriy etish (Yandex uslubida ko'p aspektli)",
    "Chevarning band kunlarini ko'rsatish"
]
for v in vazifalar:
    add_bullet_point(doc, v)

doc.add_heading("1.2 Maqsadli auditoriya", level=2)
auditoriya = [
    "Mijozlar — tikuvchi xizmatlaridan foydalanmoqchi bo'lganlar",
    "Chevarlar — o'z xizmatlarini taklif qilmoqchi bo'lgan tikuvchilar",
]
for a in auditoriya:
    add_bullet_point(doc, a)

# ==================== 2. LOYIHA MAQSADI ====================
doc.add_heading("2. LOYIHA MAQSADI", level=1)

p = doc.add_paragraph()
p.add_run("Loyihaning asosiy maqsadi — O'zbekiston bo'ylab mijozlar va chevarlar o'rtasidagi aloqani soddalashtirish va raqamlashtirish. Ilova orqali foydalanuvchilar bir necha bosqichda buyurtma berishlari, chevar haqida ma'lumot olishlari va xizmat sifatini baholashlari mumkin.")

doc.add_heading("2.1 Strategik maqsadlar", level=2)
strategik = [
    "6 oy ichida 500+ chevarni platformaga qo'shish",
    "10,000+ aktiv foydalanuvchiga erishish",
    "95%+ foydalanuvchi qoniqishini ta'minlash",
    "Buyurtma berish jarayonini soddalashtirish",
    "Chevarlar daromadini 20-30% oshirish"
]
for s in strategik:
    add_bullet_point(doc, s)

doc.add_heading("2.2 Texnik maqsadlar", level=2)
texnik_maqsad = [
    "Yuqori ishlash tezligi (3 soniyadan kam yuklanish vaqti)",
    "99.9% uptime (ishonchlilik)",
    "Android, iOS va Web platformalarini qo'llab-quvvatlash",
    "Real-time xabar almashish va ma'lumot sinxronizatsiyasi"
]
for t in texnik_maqsad:
    add_bullet_point(doc, t)

# ==================== 3. ANALITIK VAZIFALAR ====================
doc.add_heading("3. ANALITIK VAZIFALAR", level=1)

p = doc.add_paragraph()
p.add_run("Loyihani muvaffaqiyatli amalga oshirish uchun chuqur tahlil va tadqiqot o'tkazish zarur. Quyidagi analitik vazifalar belgilangan:")

doc.add_heading("3.1 Bozor tahlili", level=2)
bozor = [
    "O'zbekistondagi chevarlar sonini aniqlash (shahar va viloyatlar bo'yicha)",
    "Raqobatchilarni tahlil qilish",
    "Maqsadli auditoriyaning xatti-harakatlarini o'rganish",
    "Bozordagi tendentsiyalar va rivojlanish yo'nalishlarini aniqlash"
]
for b in bozor:
    add_bullet_point(doc, b)

doc.add_heading("3.2 Texnik tahlil", level=2)
texnik_tahlil = [
    "Flutter framework imkoniyatlarini baholash",
    "Firebase xizmatlari va cheklovlarini o'rganish",
    "To'lov tizimlarini integratsiya imkoniyatlari (Payme, Click)"
]
for t in texnik_tahlil:
    add_bullet_point(doc, t)

doc.add_heading("3.3 Foydalanuvchi tajribasi tahlili", level=2)
ux = [
    "User Journey mapping — foydalanuvchi yo'lini xaritalash",
    "UX/UI eng yaxshi amaliyotlarini o'rganish",
    "A/B testlash strategiyasini ishlab chiqish"
]
for u in ux:
    add_bullet_point(doc, u)

# ==================== 4. RIVOJLANISH VAZIFALARI ====================
doc.add_heading("4. RIVOJLANISH VAZIFALARI", level=1)

p = doc.add_paragraph()
p.add_run("Loyihani rivojlantirish uchun quyidagi texnik vazifalar belgilangan:")

doc.add_heading("4.1 Mijoz ilovasi (Flutter)", level=2)
mijoz_ilova = [
    "Foydalanuvchi autentifikatsiyasi (Email/parol bilan ro'yxatdan o'tish)",
    "Asosiy ekran — chevarlar ro'yxati va qidiruv",
    "Chevar tafsilotlari ekrani (xizmatlar, band kunlar, baho, chat)",
    "Buyurtma berish oqimi — xizmat tanlash",
    "Sevimli chevarlar ro'yxati",
    "Foydalanuvchi profili va sozlamalar (ism, viloyat tahrirlash)",
    "Xabarlar sahifasi — chevar bilan chat",
    "Baho berish tizimi (Yandex uslubida 5 aspekt: sifat, hajm, tezlik, muomala, narx)"
]
for m in mijoz_ilova:
    add_bullet_point(doc, m)

doc.add_heading("4.2 Chevar ilovasi (Flutter)", level=2)
chevar_ilova = [
    "Buyurtmalarni boshqarish (4 kategoriya: Yangi, Jarayonda, Tayyor, Yakunlangan)",
    "Xizmatlar boshqaruvi (qo'shish, tahrirlash, o'chirish)",
    "Kalendar — band kunlarni belgilash",
    "Xabarlar sahifasi — mijozlar bilan chat",
    "Profil sozlamalari (ism, viloyat tahrirlash)",
    "Til tanlash (O'zbekcha, Русский, English)"
]
for c in chevar_ilova:
    add_bullet_point(doc, c)

doc.add_heading("4.3 Backend (Firebase)", level=2)
backend = [
    "Firebase Authentication — foydalanuvchi autentifikatsiyasi",
    "Cloud Firestore — ma'lumotlar bazasi",
    "Firebase Storage — rasm va fayllar saqlash",
    "Real-time chat tizimi"
]
for b in backend:
    add_bullet_point(doc, b)

# ==================== 5. TASHKILIY VAZIFALAR ====================
doc.add_heading("5. TASHKILIY VAZIFALAR", level=1)

doc.add_heading("5.1 Jamoa tarkibi", level=2)
jamoa = [
    "Loyiha menejeri — 1 kishi",
    "Flutter dasturchilari — 2 kishi",
    "UI/UX dizayner — 1 kishi",
    "Backend dasturchisi — 1 kishi",
    "QA muhandisi — 1 kishi"
]
for j in jamoa:
    add_bullet_point(doc, j)

doc.add_heading("5.2 Aloqa va hisobot", level=2)
aloqa = [
    "Kundalik stand-up yig'ilishlar (15 daqiqa)",
    "Haftalik sprint review",
    "Oylik progress hisoboti",
    "Telegram guruhida operativ aloqa"
]
for a in aloqa:
    add_bullet_point(doc, a)

doc.add_heading("5.3 Hujjatlashtirish", level=2)
hujjat = [
    "API hujjatlari",
    "Kod hujjatlari (DartDoc)",
    "Foydalanuvchi qo'llanmasi"
]
for h in hujjat:
    add_bullet_point(doc, h)

# ==================== 6. SINOV VA SIFAT NAZORATI ====================
doc.add_heading("6. SINOV VA SIFAT NAZORATI", level=1)

p = doc.add_paragraph()
p.add_run("Mahsulot sifatini ta'minlash uchun keng qamrovli sinov strategiyasi ishlab chiqilgan:")

doc.add_heading("6.1 Sinov turlari", level=2)
sinov = [
    "Unit testlar — har bir komponent alohida sinovdan o'tkaziladi",
    "Widget testlar — UI elementlarini sinash",
    "Integration testlar — komponentlar integratsiyasini sinash",
    "End-to-end testlar — to'liq foydalanuvchi oqimini sinash",
    "Performance testlar — tezlik va yuklanish sinovlari"
]
for s in sinov:
    add_bullet_point(doc, s)

doc.add_heading("6.2 Qabul qilish mezonlari", level=2)
mezon = [
    "Barcha unit testlar muvaffaqiyatli o'tishi kerak (100%)",
    "Kod qamrovi (code coverage) — kamida 80%",
    "Kritik xatolarning yo'qligi",
    "UI/UX dizayn standartlariga muvofiqligi"
]
for m in mezon:
    add_bullet_point(doc, m)

doc.add_heading("6.3 Beta-test", level=2)
beta = [
    "50-100 beta-tester tanlash",
    "2 hafta beta-test davri",
    "Fikr-mulohazalarni yig'ish va tahlil qilish",
    "Aniqlangan xatolarni tuzatish"
]
for b in beta:
    add_bullet_point(doc, b)

# ==================== 7. FUNKSIONAL TALABLAR ====================
doc.add_heading("7. FUNKSIONAL TALABLAR", level=1)

doc.add_heading("7.1 Mijoz ilovasi", level=2)
mijoz_funk = [
    "Ro'yxatdan o'tish va kirish (Email/parol)",
    "Chevarlarni qidirish (ism, viloyat bo'yicha)",
    "Chevar tafsilotlarini ko'rish (xizmatlar, band kunlar, baholar)",
    "Buyurtma berish",
    "Buyurtmalar tarixini ko'rish",
    "Sevimli chevarlar ro'yxati",
    "Sharh va baho qoldirish (5 aspektli: sifat, hajm, tezlik, muomala, narx)",
    "Chat orqali chevar bilan bog'lanish",
    "Profil sozlamalari (ism, viloyat tahrirlash)",
    "Til tanlash (O'zbekcha, Русский, English)",
    "Chiqish (logout) tasdiqlash dialogi"
]
for m in mijoz_funk:
    add_bullet_point(doc, m)

doc.add_heading("7.2 Chevar ilovasi", level=2)
chevar_funk = [
    "Ro'yxatdan o'tish va kirish",
    "Buyurtmalarni boshqarish (4 kategoriya)",
    "Xizmatlar boshqaruvi (turi, narx diapazoni, o'rtacha vaqt)",
    "Kalendar — band kunlarni belgilash",
    "Xabarlar sahifasi — mijozlar bilan chat",
    "Profil sozlamalari (ism, viloyat tahrirlash)",
    "Til tanlash",
    "Chiqish tasdiqlash dialogi"
]
for c in chevar_funk:
    add_bullet_point(doc, c)

# ==================== 8. TEXNIK TALABLAR ====================
doc.add_heading("8. TEXNIK TALABLAR", level=1)

doc.add_heading("8.1 Texnologiya steki", level=2)

# Jadval yaratish
table = doc.add_table(rows=6, cols=2)
table.style = 'Table Grid'

# Jadval ma'lumotlari
data = [
    ("Komponent", "Texnologiya"),
    ("Frontend", "Flutter (Dart)"),
    ("Backend", "Firebase (Firestore, Auth, Storage)"),
    ("Ma'lumotlar bazasi", "Cloud Firestore"),
    ("Autentifikatsiya", "Firebase Authentication"),
    ("Xabar almashish", "Cloud Firestore Real-time"
)]

for i, (col1, col2) in enumerate(data):
    table.rows[i].cells[0].text = col1
    table.rows[i].cells[1].text = col2

doc.add_paragraph()

doc.add_heading("8.2 Tizim talablari", level=2)
tizim = [
    "Android: API 21+ (Android 5.0 Lollipop yoki undan yuqori)",
    "iOS: iOS 12.0 yoki undan yuqori",
    "Web: Barcha zamonaviy brauzerlar (Chrome, Firefox, Safari, Edge)",
    "Internet ulanishi talab qilinadi",
    "Xotira: kamida 100 MB bo'sh joy"
]
for t in tizim:
    add_bullet_point(doc, t)

# ==================== 9. DIZAYN VA INTERFEYS ====================
doc.add_heading("9. DIZAYN VA INTERFEYS", level=1)

doc.add_heading("9.1 Dizayn tamoyillari", level=2)
dizayn = [
    "Zamonaviy va minimalist dizayn",
    "Material Design 3 standartlari",
    "Adaptiv dizayn (mobil, planshet, veb)",
    "Dark/Light tema qo'llab-quvvatlash",
    "Yumshoq animatsiyalar",
    "Intuitiv navigatsiya"
]
for d in dizayn:
    add_bullet_point(doc, d)

doc.add_heading("9.2 Ranglar sxemasi", level=2)
p = doc.add_paragraph()
p.add_run("Asosiy ranglar:\n").bold = True
p.add_run("• Asosiy (Primary): #E91E63 (Pink)\n")
p.add_run("• Fon (Light): #FFFFFF\n")
p.add_run("• Fon (Dark): #121212\n")
p.add_run("• Xato: #EF4444 (Qizil)\n")
p.add_run("• Muvaffaqiyat: #10B981 (Yashil)")

doc.add_heading("9.3 Shrift va tipografiya", level=2)
tipografiya = [
    "Asosiy shrift: Google Fonts (Lato)",
    "Sarlavha: 24-28pt, Bold",
    "Asosiy matn: 14-16pt, Regular",
    "Kichik matn: 12pt, Regular"
]
for t in tipografiya:
    add_bullet_point(doc, t)

# ==================== 10. XAVFSIZLIK ====================
doc.add_heading("10. XAVFSIZLIK", level=1)

doc.add_heading("10.1 Ma'lumotlar xavfsizligi", level=2)
xavfsizlik = [
    "SSL/TLS shifrlash (HTTPS)",
    "Firebase Security Rules — ma'lumotlarga kirish nazorati",
    "Parollarni hash qilish",
    "Foydalanuvchi ma'lumotlarini shifrlash"
]
for x in xavfsizlik:
    add_bullet_point(doc, x)

doc.add_heading("10.2 Autentifikatsiya xavfsizligi", level=2)
auth_xavfsizlik = [
    "Email/parol autentifikatsiyasi",
    "Token asosida sessiya boshqaruvi",
    "Avtomatik logout funksiyasi"
]
for a in auth_xavfsizlik:
    add_bullet_point(doc, a)

doc.add_heading("10.3 Hujumlardan himoya", level=2)
himoya = [
    "XSS (Cross-Site Scripting) himoyasi",
    "CSRF (Cross-Site Request Forgery) himoyasi",
    "Firebase o'rnatilgan xavfsizlik mexanizmlari"
]
for h in himoya:
    add_bullet_point(doc, h)

# ==================== 11. MUDDATLAR VA BOSQICHLAR ====================
doc.add_heading("11. MUDDATLAR VA BOSQICHLAR", level=1)

# Bosqichlar jadvali
table2 = doc.add_table(rows=6, cols=3)
table2.style = 'Table Grid'

bosqichlar = [
    ("Bosqich", "Vazifa", "Muddat"),
    ("1-bosqich", "Tahlil va rejalashtirish", "1 hafta"),
    ("2-bosqich", "Dizayn va prototip", "2 hafta"),
    ("3-bosqich", "Asosiy funksional ishlab chiqish", "4 hafta"),
    ("4-bosqich", "Sinov va xatolarni tuzatish", "2 hafta"),
    ("5-bosqich", "Reliz va qo'llab-quvvatlash", "1 hafta")
]

for i, (col1, col2, col3) in enumerate(bosqichlar):
    table2.rows[i].cells[0].text = col1
    table2.rows[i].cells[1].text = col2
    table2.rows[i].cells[2].text = col3

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Umumiy loyiha muddati: 10 hafta (taxminan 2.5 oy)").bold = True

# ==================== XULOSA ====================
doc.add_page_break()
doc.add_heading("XULOSA", level=1)

p = doc.add_paragraph()
p.add_run("TailorMatch").bold = True
p.add_run(" loyihasi O'zbekiston bozoriga yo'naltirilgan innovatsion mahsulot bo'lib, tikuvchilik sohasini raqamlashtirish va chevarlarga yangi mijozlar oqimini ta'minlash maqsadini ko'zlaydi.")

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Ushbu texnik topshiriq loyihaning barcha jihatlarini qamrab olgan bo'lib, jamoa a'zolari uchun aniq yo'l-yo'riq vazifasini bajaradi. Loyiha muvaffaqiyatli amalga oshirilgandan so'ng, O'zbekiston bo'ylab minglab foydalanuvchilarga xizmat ko'rsatishga tayyor bo'ladi.")

doc.add_paragraph()
doc.add_paragraph()

# Imzo
p = doc.add_paragraph()
p.add_run(f"Tayyorladi: Rustamjon Abduvahobov\n")
p.add_run(f"Sana: {datetime.now().strftime('%d.%m.%Y')}")

# Faylni saqlash va yuklab olish
filename = "TailorMatch_Texnik_Topshiriq.docx"
doc.save(filename)
files.download(filename)

print(f"✅ Hujjat muvaffaqiyatli yaratildi: {filename}")
print("📥 Yuklab olish avtomatik boshlanadi...")
