"""
TailorMatch - Technical Specification Generator (English Version)
Run this code in Google Colab to generate Word document
"""

# Install required libraries
!pip install python-docx -q

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from google.colab import files
from datetime import datetime

# Create document
doc = Document()

# Style functions
def set_heading_style(paragraph, size=18, bold=True, color=None):
    run = paragraph.runs[0]
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_bullet_point(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

# ==================== TITLE PAGE ====================
title = doc.add_heading('TECHNICAL SPECIFICATION', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('TAILORMATCH')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(233, 30, 99)  # Pink color

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run("Mobile Application for Connecting\nTailors and Clients in Uzbekistan")
run.font.size = Pt(14)

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run(f'Version: 1.0\n')
info.add_run(f'Date: {datetime.now().strftime("%d.%m.%Y")}\n')
info.add_run('Author: Rustamjon Abduvahobov')

doc.add_page_break()

# ==================== TABLE OF CONTENTS ====================
doc.add_heading('TABLE OF CONTENTS', level=1)
contents = [
    "1. Definition of Project Tasks and Objectives",
    "2. Project Goals",
    "3. Analytical Tasks",
    "4. Development Tasks",
    "5. Organizational Tasks",
    "6. Testing and Quality Assurance",
    "7. Functional Requirements",
    "8. Technical Requirements",
    "9. Design and Interface",
    "10. Security",
    "11. Timeline and Phases"
]
for item in contents:
    doc.add_paragraph(item)

doc.add_page_break()

# ==================== 1. PROJECT TASKS ====================
doc.add_heading("1. DEFINITION OF PROJECT TASKS AND OBJECTIVES", level=1)

p = doc.add_paragraph()
p.add_run("TailorMatch").bold = True
p.add_run(" — is a modern mobile application designed to establish connections between tailors and clients in Uzbekistan. The project is being developed with the goal of digitizing the tailoring industry and providing convenient services to customers.")

doc.add_heading("1.1 Main Tasks", level=2)
tasks = [
    "Provide users with the ability to search and browse tailors",
    "Implement an online ordering system",
    "Create an order management panel for tailors",
    "Enable effective chat communication between clients and tailors",
    "Filter tailors by region (viloyat)",
    "Implement a rating and review system (Yandex-style multi-aspect)",
    "Display tailor's busy days calendar"
]
for t in tasks:
    add_bullet_point(doc, t)

doc.add_heading("1.2 Target Audience", level=2)
audience = [
    "Clients — those who want to use tailoring services",
    "Tailors — seamstresses who want to offer their services",
]
for a in audience:
    add_bullet_point(doc, a)

# ==================== 2. PROJECT GOALS ====================
doc.add_heading("2. PROJECT GOALS", level=1)

p = doc.add_paragraph()
p.add_run("The main goal of the project is to simplify and digitize communication between clients and tailors throughout Uzbekistan. Through the application, users can place orders in a few steps, get information about tailors, and rate service quality.")

doc.add_heading("2.1 Strategic Goals", level=2)
strategic = [
    "Onboard 500+ tailors to the platform within 6 months",
    "Reach 10,000+ active users",
    "Achieve 95%+ user satisfaction",
    "Simplify the ordering process",
    "Increase tailor income by 20-30%"
]
for s in strategic:
    add_bullet_point(doc, s)

doc.add_heading("2.2 Technical Goals", level=2)
technical_goals = [
    "High performance (loading time less than 3 seconds)",
    "99.9% uptime (reliability)",
    "Support Android, iOS, and Web platforms",
    "Real-time messaging and data synchronization"
]
for t in technical_goals:
    add_bullet_point(doc, t)

# ==================== 3. ANALYTICAL TASKS ====================
doc.add_heading("3. ANALYTICAL TASKS", level=1)

p = doc.add_paragraph()
p.add_run("Deep analysis and research are necessary for successful project implementation. The following analytical tasks have been defined:")

doc.add_heading("3.1 Market Analysis", level=2)
market = [
    "Determine the number of tailors in Uzbekistan (by city and region)",
    "Analyze competitors",
    "Study target audience behavior",
    "Identify market trends and development directions"
]
for m in market:
    add_bullet_point(doc, m)

doc.add_heading("3.2 Technical Analysis", level=2)
technical_analysis = [
    "Evaluate Flutter framework capabilities",
    "Study Firebase services and limitations",
    "Payment system integration options (Payme, Click)"
]
for t in technical_analysis:
    add_bullet_point(doc, t)

doc.add_heading("3.3 User Experience Analysis", level=2)
ux = [
    "User Journey mapping",
    "Study UX/UI best practices",
    "Develop A/B testing strategy"
]
for u in ux:
    add_bullet_point(doc, u)

# ==================== 4. DEVELOPMENT TASKS ====================
doc.add_heading("4. DEVELOPMENT TASKS", level=1)

p = doc.add_paragraph()
p.add_run("The following technical tasks have been defined for project development:")

doc.add_heading("4.1 Client Application (Flutter)", level=2)
client_app = [
    "User authentication (Email/password registration)",
    "Main screen — tailor list and search",
    "Tailor detail screen (services, busy days, ratings, chat)",
    "Order placement flow — service selection",
    "Favorite tailors list",
    "User profile and settings (name, region editing)",
    "Messages page — chat with tailor",
    "Rating system (Yandex-style 5 aspects: quality, size, speed, service, price)"
]
for c in client_app:
    add_bullet_point(doc, c)

doc.add_heading("4.2 Tailor Application (Flutter)", level=2)
tailor_app = [
    "Order management (4 categories: New, In Progress, Ready, Completed)",
    "Service management (add, edit, delete)",
    "Calendar — mark busy days",
    "Messages page — chat with clients",
    "Profile settings (name, region editing)",
    "Language selection (Uzbek, Russian, English)"
]
for t in tailor_app:
    add_bullet_point(doc, t)

doc.add_heading("4.3 Backend (Firebase)", level=2)
backend = [
    "Firebase Authentication — user authentication",
    "Cloud Firestore — database",
    "Firebase Storage — image and file storage",
    "Real-time chat system"
]
for b in backend:
    add_bullet_point(doc, b)

# ==================== 5. ORGANIZATIONAL TASKS ====================
doc.add_heading("5. ORGANIZATIONAL TASKS", level=1)

doc.add_heading("5.1 Team Composition", level=2)
team = [
    "Project Manager — 1 person",
    "Flutter Developers — 2 people",
    "UI/UX Designer — 1 person",
    "Backend Developer — 1 person",
    "QA Engineer — 1 person"
]
for t in team:
    add_bullet_point(doc, t)

doc.add_heading("5.2 Communication and Reporting", level=2)
communication = [
    "Daily stand-up meetings (15 minutes)",
    "Weekly sprint review",
    "Monthly progress report",
    "Operational communication via Telegram group"
]
for c in communication:
    add_bullet_point(doc, c)

doc.add_heading("5.3 Documentation", level=2)
documentation = [
    "API documentation",
    "Code documentation (DartDoc)",
    "User manual"
]
for d in documentation:
    add_bullet_point(doc, d)

# ==================== 6. TESTING AND QUALITY ASSURANCE ====================
doc.add_heading("6. TESTING AND QUALITY ASSURANCE", level=1)

p = doc.add_paragraph()
p.add_run("A comprehensive testing strategy has been developed to ensure product quality:")

doc.add_heading("6.1 Testing Types", level=2)
testing = [
    "Unit tests — each component tested separately",
    "Widget tests — testing UI elements",
    "Integration tests — testing component integration",
    "End-to-end tests — testing complete user flows",
    "Performance tests — speed and load testing"
]
for t in testing:
    add_bullet_point(doc, t)

doc.add_heading("6.2 Acceptance Criteria", level=2)
criteria = [
    "All unit tests must pass successfully (100%)",
    "Code coverage — at least 80%",
    "Absence of critical errors",
    "Compliance with UI/UX design standards"
]
for c in criteria:
    add_bullet_point(doc, c)

doc.add_heading("6.3 Beta Testing", level=2)
beta = [
    "Select 50-100 beta testers",
    "2-week beta testing period",
    "Collect and analyze feedback",
    "Fix identified bugs"
]
for b in beta:
    add_bullet_point(doc, b)

# ==================== 7. FUNCTIONAL REQUIREMENTS ====================
doc.add_heading("7. FUNCTIONAL REQUIREMENTS", level=1)

doc.add_heading("7.1 Client Application", level=2)
client_func = [
    "Registration and login (Email/password)",
    "Search tailors (by name, region)",
    "View tailor details (services, busy days, ratings)",
    "Place orders",
    "View order history",
    "Favorite tailors list",
    "Leave reviews and ratings (5 aspects: quality, size, speed, service, price)",
    "Connect with tailor via chat",
    "Profile settings (name, region editing)",
    "Language selection (Uzbek, Russian, English)",
    "Logout confirmation dialog"
]
for c in client_func:
    add_bullet_point(doc, c)

doc.add_heading("7.2 Tailor Application", level=2)
tailor_func = [
    "Registration and login",
    "Order management (4 categories)",
    "Service management (type, price range, average time)",
    "Calendar — mark busy days",
    "Messages page — chat with clients",
    "Profile settings (name, region editing)",
    "Language selection",
    "Logout confirmation dialog"
]
for t in tailor_func:
    add_bullet_point(doc, t)

# ==================== 8. TECHNICAL REQUIREMENTS ====================
doc.add_heading("8. TECHNICAL REQUIREMENTS", level=1)

doc.add_heading("8.1 Technology Stack", level=2)

# Create table
table = doc.add_table(rows=6, cols=2)
table.style = 'Table Grid'

# Table data
data = [
    ("Component", "Technology"),
    ("Frontend", "Flutter (Dart)"),
    ("Backend", "Firebase (Firestore, Auth, Storage)"),
    ("Database", "Cloud Firestore"),
    ("Authentication", "Firebase Authentication"),
    ("Messaging", "Cloud Firestore Real-time")
]

for i, (col1, col2) in enumerate(data):
    table.rows[i].cells[0].text = col1
    table.rows[i].cells[1].text = col2

doc.add_paragraph()

doc.add_heading("8.2 System Requirements", level=2)
system = [
    "Android: API 21+ (Android 5.0 Lollipop or higher)",
    "iOS: iOS 12.0 or higher",
    "Web: All modern browsers (Chrome, Firefox, Safari, Edge)",
    "Internet connection required",
    "Storage: minimum 100 MB free space"
]
for s in system:
    add_bullet_point(doc, s)

# ==================== 9. DESIGN AND INTERFACE ====================
doc.add_heading("9. DESIGN AND INTERFACE", level=1)

doc.add_heading("9.1 Design Principles", level=2)
design = [
    "Modern and minimalist design",
    "Material Design 3 standards",
    "Adaptive design (mobile, tablet, web)",
    "Dark/Light theme support",
    "Smooth animations",
    "Intuitive navigation"
]
for d in design:
    add_bullet_point(doc, d)

doc.add_heading("9.2 Color Scheme", level=2)
p = doc.add_paragraph()
p.add_run("Primary Colors:\n").bold = True
p.add_run("• Primary: #E91E63 (Pink)\n")
p.add_run("• Background (Light): #FFFFFF\n")
p.add_run("• Background (Dark): #121212\n")
p.add_run("• Error: #EF4444 (Red)\n")
p.add_run("• Success: #10B981 (Green)")

doc.add_heading("9.3 Typography", level=2)
typography = [
    "Primary font: Google Fonts (Lato)",
    "Headings: 24-28pt, Bold",
    "Body text: 14-16pt, Regular",
    "Small text: 12pt, Regular"
]
for t in typography:
    add_bullet_point(doc, t)

# ==================== 10. SECURITY ====================
doc.add_heading("10. SECURITY", level=1)

doc.add_heading("10.1 Data Security", level=2)
security = [
    "SSL/TLS encryption (HTTPS)",
    "Firebase Security Rules — data access control",
    "Password hashing",
    "User data encryption"
]
for s in security:
    add_bullet_point(doc, s)

doc.add_heading("10.2 Authentication Security", level=2)
auth_security = [
    "Email/password authentication",
    "Token-based session management",
    "Automatic logout functionality"
]
for a in auth_security:
    add_bullet_point(doc, a)

doc.add_heading("10.3 Attack Protection", level=2)
protection = [
    "XSS (Cross-Site Scripting) protection",
    "CSRF (Cross-Site Request Forgery) protection",
    "Firebase built-in security mechanisms"
]
for p_item in protection:
    add_bullet_point(doc, p_item)

# ==================== 11. TIMELINE AND PHASES ====================
doc.add_heading("11. TIMELINE AND PHASES", level=1)

# Phases table
table2 = doc.add_table(rows=6, cols=3)
table2.style = 'Table Grid'

phases = [
    ("Phase", "Task", "Duration"),
    ("Phase 1", "Analysis and planning", "1 week"),
    ("Phase 2", "Design and prototype", "2 weeks"),
    ("Phase 3", "Core functionality development", "4 weeks"),
    ("Phase 4", "Testing and bug fixes", "2 weeks"),
    ("Phase 5", "Release and support", "1 week")
]

for i, (col1, col2, col3) in enumerate(phases):
    table2.rows[i].cells[0].text = col1
    table2.rows[i].cells[1].text = col2
    table2.rows[i].cells[2].text = col3

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Total project duration: 10 weeks (approximately 2.5 months)").bold = True

# ==================== CONCLUSION ====================
doc.add_page_break()
doc.add_heading("CONCLUSION", level=1)

p = doc.add_paragraph()
p.add_run("TailorMatch").bold = True
p.add_run(" is an innovative product targeted at the Uzbekistan market, aiming to digitize the tailoring industry and provide tailors with a new stream of customers.")

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("This technical specification comprehensively covers all aspects of the project and serves as a clear guideline for team members. Upon successful implementation, the project will be ready to serve thousands of users across Uzbekistan.")

doc.add_paragraph()
doc.add_paragraph()

# Signature
p = doc.add_paragraph()
p.add_run(f"Prepared by: Rustamjon Abduvahobov\n")
p.add_run(f"Date: {datetime.now().strftime('%d.%m.%Y')}")

# Save and download file
filename = "TailorMatch_Technical_Specification_EN.docx"
doc.save(filename)
files.download(filename)

print(f"✅ Document created successfully: {filename}")
print("📥 Download will start automatically...")
